from docx import Document
from docx.shared import Inches

document =Document()
# document.add_heading("Multiple Choice Questions")

# document.add_paragraph("blah blah blah")
table =document.add_table(rows=63,cols=16)
table.style='TableGrid'

# cell.text='parrot, clever bird'
p=29
i=0
# print((table))
for cell in table.columns[0].cells: # this is to set width of columns
	cell.width=Inches(0.3)
p=0
for k in range(63): 	# to traverse rows
	for i in range(0,16,2): #to traverse columns
		cell=table.cell(k,i)
		# font=run.font
		# font.size=pt(10)
		cell.text="%s" %(1+p)
		p+=1
		if(p==500):
			break
	if(p==500):
		break			
		# print(k) 
	# i+=2
#this was created to use full page or if different set of questions are required	
# print("TABLE2 \n\n")
# table2 =document.add_table(rows=22,cols=16)	
# table2.style='TableGrid'
# p=0
# for k in range(22):
# 	for i in range(0,16,2):
# 		cell=table2.cell(k,i)
# 		# font=run.font
# 		# font.size=pt(10)
# 		cell.text="%s" %(1+p)
# 		p+=1
# 		if(p==175):
# 			break
# 	if(p==175):
# 		break	

document.save('../objective_solving_sheet/sheet.docx')
# give path of your directory(by doing pwd you will get path of your director) or you can use the same given above